"""Unstructured document loader with multimodal support.

支持从文档中提取文本和图片，为多模态嵌入模型 (qwen3-vl-embedding-8b) 提供数据支持。
"""
import base64
import hashlib
import logging
import os
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

from ...utils.image_storage import get_image_storage_manager

logger = logging.getLogger(__name__)


def _extract_docx_images_with_position(file_path: str, figures_dir: Path, doc_id: str,
                                        save_images: bool = True, embed_base64: bool = True) -> Tuple[List[Dict[str, Any]], Dict[int, List[int]]]:
    """使用 python-docx 从 DOCX 文件中提取图片，并精确记录其在段落中的位置。
    
    通过遍历段落和 run，检查每个 run 中是否有 <a:blip> 元素来定位图片位置，
    而不是简单地遍历 relationship 列表。
    
    使用公共图片存储策略：大图保存文件，小图内联 base64。
    
    Args:
        file_path: DOCX 文件路径
        figures_dir: 图片保存目录
        doc_id: 文档 ID
        save_images: 是否保存图片文件（已废弃，由存储策略自动决定）
        embed_base64: 是否嵌入 base64 数据（已废弃，由存储策略自动决定）
    
    Returns:
        Tuple[图片信息列表, 段落-图片索引映射 {段落索引: [图片索引列表]}]
    """
    images = []
    # 记录每个段落包含哪些图片的索引
    paragraph_image_map: Dict[int, List[int]] = {}
    
    # 使用公共图片存储管理器
    storage_manager = get_image_storage_manager()
    
    try:
        from docx import Document
        
        # XML 命名空间 URI 定义 (用于 findall)
        PIC_NS = '{http://schemas.openxmlformats.org/drawingml/2006/picture}'
        BLIP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
        REL_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
        WP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
        
        doc = Document(file_path)
        paragraphs = list(doc.paragraphs)
        image_index = 0
        
        # 遍历每个段落
        for para_idx, paragraph in enumerate(paragraphs):
            # 获取段落前后的文本作为上下文
            context_before = paragraphs[para_idx - 1].text if para_idx > 0 else None
            context_after = paragraphs[para_idx + 1].text if para_idx < len(paragraphs) - 1 else None
            
            # 查找段落中的所有图片 (pic:pic 元素) - 使用 findall 和完整命名空间
            pics = paragraph._element.findall(f'.//{PIC_NS}pic')
            
            for pic in pics:
                # 获取图片的 embed 引用 - 查找 a:blip 元素
                blips = pic.findall(f'.//{BLIP_NS}blip')
                if not blips:
                    continue
                
                # 获取 r:embed 属性
                embed_id = blips[0].get(f'{REL_NS}embed')
                if not embed_id:
                    continue
                
                try:
                    # 通过 embed_id 获取图片数据
                    image_part = doc.part.related_parts.get(embed_id)
                    if not image_part:
                        continue
                    
                    image_bytes = image_part.blob
                    content_type = image_part.content_type
                    
                    # 尝试获取 alt text (描述文本) - 从 docPr 元素
                    alt_text = None
                    try:
                        # 向上查找 wp:inline 或 wp:anchor，然后找 docPr
                        parent = pic.getparent()
                        while parent is not None:
                            tag = parent.tag
                            if tag in (f'{WP_NS}inline', f'{WP_NS}anchor'):
                                # 找到 docPr 元素
                                doc_pr = parent.find(f'{WP_NS}docPr')
                                if doc_pr is not None:
                                    alt_text = doc_pr.get('descr') or doc_pr.get('name')
                                break
                            parent = parent.getparent()
                    except Exception:
                        pass
                    
                    # 使用公共存储策略处理图片
                    storage_result = storage_manager.process_image(
                        image_bytes=image_bytes,
                        image_index=image_index,
                        figures_dir=figures_dir,
                        doc_id=doc_id,
                        page_number=1,  # DOCX 不分页，统一设为 1
                        mime_type=content_type,
                    )
                    
                    image_info = {
                        "page_number": 1,  # DOCX 不分页，统一设为 1
                        "image_index": image_index,
                        "paragraph_index": para_idx,  # 记录段落位置（关键改进）
                        "caption": None,
                        "alt_text": alt_text,
                        "bbox": None,
                        "file_path": storage_result["file_path"],
                        "base64_data": storage_result["base64_data"],
                        "thumbnail_base64": storage_result.get("thumbnail_base64"),
                        "mime_type": storage_result["mime_type"],
                        "width": storage_result["width"],
                        "height": storage_result["height"],
                        "original_size": storage_result["original_size"],
                        "context_before": context_before[:200] if context_before else None,
                        "context_after": context_after[:200] if context_after else None,
                        "image_type": "embedded",
                        "ocr_text": None,
                        "storage_type": storage_result.get("storage_type"),
                    }
                    
                    images.append(image_info)
                    
                    # 记录段落-图片映射
                    if para_idx not in paragraph_image_map:
                        paragraph_image_map[para_idx] = []
                    paragraph_image_map[para_idx].append(image_index)
                    
                    image_index += 1
                    
                except Exception as e:
                    logger.warning(f"[python-docx] Failed to extract image from paragraph {para_idx}: {e}")
                    continue
        
        logger.info(f"[python-docx] Extracted {len(images)} images with position from DOCX, "
                    f"distributed across {len(paragraph_image_map)} paragraphs")
        
    except ImportError:
        logger.warning("[python-docx] python-docx not installed, cannot extract DOCX images")
    except Exception as e:
        logger.warning(f"[python-docx] Failed to extract images from DOCX: {e}")
    
    return images, paragraph_image_map


def _extract_docx_paragraphs(file_path: str) -> List[str]:
    """使用 python-docx 提取 DOCX 文件的所有段落文本。
    
    Args:
        file_path: DOCX 文件路径
    
    Returns:
        段落文本列表
    """
    try:
        from docx import Document
        doc = Document(file_path)
        return [p.text for p in doc.paragraphs]
    except Exception as e:
        logger.warning(f"[python-docx] Failed to extract paragraphs: {e}")
        return []


def _iter_docx_block_items(doc):
    """按文档原始顺序迭代 DOCX 中的段落和表格。
    
    这是业内标准做法，通过遍历 Word 的底层 XML 结构来保持元素的原始顺序。
    
    Args:
        doc: python-docx Document 对象
        
    Yields:
        Tuple[str, Any]: ('paragraph', Paragraph, para_idx) 或 ('table', Table, table_idx)
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    
    para_idx = 0
    table_idx = 0
    
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield ('paragraph', Paragraph(child, doc), para_idx)
            para_idx += 1
        elif isinstance(child, CT_Tbl):
            yield ('table', Table(child, doc), table_idx)
            table_idx += 1


class UnstructuredLoader:
    """Load documents using Unstructured library with multimodal support.
    
    支持:
    - 文本提取 (所有支持格式)
    - 图片提取和保存 (PDF, DOCX 等)
    - 表格结构识别
    - OCR 文字识别
    """
    
    # 默认图片保存目录 (相对于工作目录)
    DEFAULT_FIGURES_DIR = "figures"
    
    def __init__(self, figures_dir: Optional[str] = None):
        """Initialize Unstructured loader.
        
        Args:
            figures_dir: 图片保存目录，默认为 'figures'
        """
        self.name = "unstructured"
        self._unavailable_reason = None
        self.figures_dir = figures_dir or self.DEFAULT_FIGURES_DIR
    
    def is_available(self) -> bool:
        """Check if Unstructured library is available."""
        try:
            from unstructured.partition.auto import partition
            self._unavailable_reason = None
            return True
        except ImportError as e:
            self._unavailable_reason = f"Unstructured library not installed: {str(e)}"
            return False
    
    def get_unavailable_reason(self) -> str:
        """Get reason why loader is unavailable."""
        return self._unavailable_reason or "Unknown reason"
    
    def extract_text(
        self, 
        file_path: str,
        extract_images: bool = True,
        save_images: bool = True,
        embed_images_base64: bool = True  # 默认启用 base64，用于多模态嵌入
    ) -> Dict[str, Any]:
        """Extract text and images from document using Unstructured.
        
        Args:
            file_path: Path to document file
            extract_images: 是否提取图片信息
            save_images: 是否将图片保存为文件
            embed_images_base64: 是否将图片嵌入为 Base64 (用于多模态处理)
            
        Returns:
            Dictionary with extracted content and metadata, including:
            - pages: 分页文本内容
            - images: 图片信息列表 (支持多模态)
            - tables: 表格信息列表
        """
        try:
            from unstructured.partition.auto import partition
            from unstructured.documents.elements import Image as UnstructuredImage
            
            path = Path(file_path)
            doc_id = self._generate_doc_id(path)
            file_ext = path.suffix.lower()
            
            # 确保图片目录存在
            figures_path = Path(self.figures_dir)
            if save_images and not figures_path.exists():
                figures_path.mkdir(parents=True, exist_ok=True)
            
            # 对于 DOCX 文件，使用改进的 python-docx 方法提取图片（包含位置信息）
            docx_images = []
            docx_paragraph_image_map: Dict[int, List[int]] = {}
            docx_paragraphs: List[str] = []
            is_docx = file_ext in (".docx", ".doc")
            
            if is_docx and extract_images:
                docx_images, docx_paragraph_image_map = _extract_docx_images_with_position(
                    file_path=file_path,
                    figures_dir=figures_path,
                    doc_id=doc_id,
                    save_images=save_images,
                    embed_base64=embed_images_base64
                )
                # 同时提取段落文本用于位置对应
                docx_paragraphs = _extract_docx_paragraphs(file_path)
                logger.info(f"[Unstructured] Extracted {len(docx_images)} images from DOCX with position info, "
                            f"{len(docx_paragraphs)} paragraphs")
            
            # 构建 partition 参数 (根据文件类型差异化配置)
            partition_kwargs = {
                "filename": file_path,
                "strategy": "hi_res",
                "infer_table_structure": True,
                "languages": ["chi_sim", "eng"],  # 使用新参数替代 ocr_languages
            }
            
            # PDF 专用图片提取参数
            if file_ext == ".pdf" and extract_images:
                partition_kwargs["extract_images_in_pdf"] = True
                partition_kwargs["extract_image_block_to_payload"] = True
                partition_kwargs["extract_image_block_types"] = ["Image", "Figure"]
            
            # PPTX 等 Office 文档：指定图片输出目录 (DOCX 使用 python-docx 处理)
            if file_ext in (".pptx", ".ppt") and extract_images:
                # 为当前文档创建专属图片目录
                doc_figures_dir = figures_path / doc_id
                doc_figures_dir.mkdir(parents=True, exist_ok=True)
                partition_kwargs["extract_image_block_output_dir"] = str(doc_figures_dir)
                partition_kwargs["extract_image_block_types"] = ["Image", "Figure"]
            
            elements = partition(**partition_kwargs)
            
            # Process elements
            pages_text: Dict[int, List[str]] = {}
            full_text: List[str] = []
            images: List[Dict[str, Any]] = []
            tables: List[Dict[str, Any]] = []
            
            # 统计元素类型用于调试
            element_type_counts = {}
            
            # 用于收集 xlsx 文件的 sheet 名称 (page_number -> sheet_name 映射)
            page_to_sheet_name: Dict[int, str] = {}
            is_xlsx = file_ext in ('.xlsx', '.xls')
            
            # 对于 DOCX 文件：使用 iter_block_items 按原文档顺序遍历段落和表格
            if is_docx:
                # 图片已有完整信息，直接添加
                if docx_images:
                    images.extend(docx_images)
                
                # 预处理 unstructured 的表格元素（用于获取高质量的表格解析结果）
                unstructured_tables = []
                for element in elements:
                    element_type = type(element).__name__
                    element_type_counts[element_type] = element_type_counts.get(element_type, 0) + 1
                    if element_type == 'Table':
                        page_num = 1
                        table_info = self._process_table_element(element, page_num, len(unstructured_tables))
                        if table_info:
                            unstructured_tables.append(table_info)
                
                # 使用 iter_block_items 按原文档顺序遍历
                from docx import Document
                doc = Document(file_path)
                
                page_num = 1
                pages_text[page_num] = []
                table_idx = 0
                
                for block_type, block, idx in _iter_docx_block_items(doc):
                    if block_type == 'paragraph':
                        para_text = block.text
                        # 添加段落文本
                        if para_text:
                            pages_text[page_num].append(para_text)
                            full_text.append(para_text)
                        
                        # 检查该段落是否有图片，如果有则在文本后添加图片占位符
                        if idx in docx_paragraph_image_map:
                            for img_idx in docx_paragraph_image_map[idx]:
                                if img_idx < len(docx_images):
                                    img_info = docx_images[img_idx]
                                    # 占位符索引从1开始（人类可读），images数组的image_index从0开始
                                    placeholder = f"[IMAGE_{img_idx + 1}: {img_info.get('alt_text') or '图片'}]"
                                    pages_text[page_num].append(placeholder)
                                    full_text.append(placeholder)
                    
                    elif block_type == 'table':
                        # 在正确位置插入表格
                        # 优先使用 unstructured 解析的表格（质量更高）
                        if table_idx < len(unstructured_tables):
                            table_info = unstructured_tables[table_idx]
                            tables.append(table_info)
                            table_text = table_info.get("markdown") or ""
                        else:
                            # 降级使用 python-docx 解析
                            table_text = self._python_docx_table_to_markdown(block)
                            table_info = {
                                "page_number": page_num,
                                "table_index": len(tables),
                                "headers": [],
                                "rows": [],
                                "caption": None,
                                "markdown": table_text,
                                "html": None
                            }
                            tables.append(table_info)
                        
                        if table_text:
                            pages_text[page_num].append(table_text)
                            full_text.append(table_text)
                        
                        table_idx += 1
                
                logger.info(f"[Unstructured] DOCX processed with position-aware images and tables")
            
            else:
                # 非 DOCX 文件：使用原有逻辑处理
                image_index = 0
                prev_text = ""  # 跟踪上一个文本元素，用于图片上下文
                
                for i, element in enumerate(elements):
                    raw_page_num = getattr(element.metadata, 'page_number', None) if hasattr(element, 'metadata') else None
                    page_num = raw_page_num if isinstance(raw_page_num, int) and raw_page_num > 0 else 1
                    element_type = type(element).__name__
                    element_type_counts[element_type] = element_type_counts.get(element_type, 0) + 1
                    
                    # 对于 xlsx 文件，提取 page_name（即 sheet 名称）
                    if is_xlsx and hasattr(element, 'metadata'):
                        page_name = getattr(element.metadata, 'page_name', None)
                        if page_name and page_num not in page_to_sheet_name:
                            page_to_sheet_name[page_num] = page_name
                    
                    if page_num not in pages_text:
                        pages_text[page_num] = []
                    
                    # 处理图片元素
                    if element_type in ('Image', 'Figure') and extract_images:
                        image_info = self._process_image_element(
                            element=element,
                            image_index=image_index,
                            page_num=page_num,
                            doc_id=doc_id,
                            figures_path=figures_path,
                            save_images=save_images,
                            embed_base64=embed_images_base64,
                            context_before=prev_text[-200:] if prev_text else None  # 保留最后200字符作为上下文
                        )
                        if image_info:
                            images.append(image_info)
                            # 在文本中插入图片占位符
                            # 占位符索引从1开始（人类可读），images数组的image_index从0开始
                            placeholder = f"[IMAGE_{image_index + 1}: {image_info.get('caption') or '图片'}]"
                            pages_text[page_num].append(placeholder)
                            full_text.append(placeholder)
                            image_index += 1
                    
                    # 处理表格元素
                    elif element_type == 'Table':
                        table_info = self._process_table_element(element, page_num, len(tables))
                        if table_info:
                            tables.append(table_info)
                            # 使用 Markdown 格式的表格文本（如果有），否则降级使用纯文本
                            table_text = table_info.get("markdown") or element.text
                            pages_text[page_num].append(table_text)
                            full_text.append(table_text)
                            prev_text = table_text
                    
                    # 处理普通文本元素
                    else:
                        text = element.text
                        if text:
                            pages_text[page_num].append(text)
                            full_text.append(text)
                            prev_text = text
                            
                            # 更新上一个图片的 context_after
                            if images and not images[-1].get('context_after'):
                                images[-1]['context_after'] = text[:200]  # 前200字符
            
            # Format pages
            pages = []
            for page_num in sorted(pages_text.keys()):
                text = "\n".join(pages_text[page_num])
                page_info = {
                    "page_number": page_num,
                    "text": text,
                    "char_count": len(text)
                }
                # 对于 xlsx 文件，添加 sheet_name
                if is_xlsx and page_num in page_to_sheet_name:
                    page_info["sheet_name"] = page_to_sheet_name[page_num]
                pages.append(page_info)
            
            # 日志：元素类型统计
            logger.info(f"[Unstructured] Element types: {element_type_counts}")
            
            # 构建元数据
            result_metadata = {
                "elements_count": len(elements),
                "image_count": len(images),
                "table_count": len(tables),
                "multimodal_ready": len(images) > 0,  # 标记是否有多模态内容
                "element_types": element_type_counts  # 添加元素类型统计
            }
            
            # 对于 xlsx 文件，添加 sheet 相关元数据（与 docling 输出保持一致）
            if is_xlsx and page_to_sheet_name:
                # 按页码排序获取 sheet 名称列表
                sheet_names = [page_to_sheet_name[p] for p in sorted(page_to_sheet_name.keys())]
                result_metadata["sheet_count"] = len(sheet_names)
                result_metadata["sheet_names"] = sheet_names
            
            return {
                "success": True,
                "loader": self.name,
                "metadata": result_metadata,
                "pages": pages,
                "images": images,
                "tables": tables,
                "full_text": "\n\n".join(full_text),
                "total_pages": len(pages),
                "total_chars": sum(p["char_count"] for p in pages)
            }
            
        except ImportError as e:
            return {
                "success": False,
                "loader": self.name,
                "error": f"Unstructured dependency missing: {str(e)}"
            }
        except Exception as e:
            logger.error(f"Document extraction failed: {e}", exc_info=True)
            return {
                "success": False,
                "loader": self.name,
                "error": str(e)
            }
    
    def _process_image_element(
        self,
        element,
        image_index: int,
        page_num: int,
        doc_id: str,
        figures_path: Path,
        save_images: bool,
        embed_base64: bool,
        context_before: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """Process an image element and extract multimodal data.
        
        Args:
            element: Unstructured Image/Figure element
            image_index: 图片序号
            page_num: 页码
            doc_id: 文档 ID
            figures_path: 图片保存目录
            save_images: 是否保存图片文件
            embed_base64: 是否嵌入 Base64
            context_before: 图片前的文本上下文
        
        Returns:
            图片信息字典，支持多模态处理
        """
        try:
            metadata = element.metadata if hasattr(element, 'metadata') else None
            
            # 基础信息
            image_info = {
                "page_number": page_num,
                "image_index": image_index,
                "caption": None,
                "alt_text": None,
                "bbox": None,
                "file_path": None,
                "base64_data": None,
                "mime_type": None,
                "width": None,
                "height": None,
                "context_before": context_before,
                "context_after": None,
                "image_type": "figure",
                "ocr_text": None
            }
            
            # 提取元数据
            if metadata:
                # 标题/描述
                image_info["caption"] = getattr(metadata, 'image_caption', None)
                image_info["alt_text"] = getattr(metadata, 'alt_text', None) or element.text
                
                # 位置信息
                coords = getattr(metadata, 'coordinates', None)
                if coords:
                    image_info["bbox"] = {
                        "x1": coords.points[0][0] if coords.points else 0,
                        "y1": coords.points[0][1] if coords.points else 0,
                        "x2": coords.points[2][0] if len(coords.points) > 2 else 0,
                        "y2": coords.points[2][1] if len(coords.points) > 2 else 0
                    }
                
                # 方式1: 从 image_path 获取已保存的图片 (DOCX/PPTX 等)
                image_path = getattr(metadata, 'image_path', None)
                if image_path and os.path.exists(image_path):
                    image_info["file_path"] = image_path
                    # 读取图片获取 base64 和尺寸
                    try:
                        with open(image_path, 'rb') as f:
                            image_data = f.read()
                        
                        # 检测 MIME 类型
                        image_info["mime_type"] = self._detect_image_mime_from_bytes(image_data)
                        
                        # 获取尺寸
                        size = self._get_image_size(image_data)
                        if size:
                            image_info["width"], image_info["height"] = size
                        
                        # 嵌入 Base64
                        if embed_base64:
                            image_info["base64_data"] = base64.b64encode(image_data).decode('utf-8')
                        
                        logger.debug(f"Loaded image from path: {image_path}")
                    except Exception as e:
                        logger.warning(f"Failed to read image from path {image_path}: {e}")
                
                # 方式2: 从 base64 payload 获取 (PDF)
                if not image_info["file_path"]:
                    image_base64 = getattr(metadata, 'image_base64', None)
                    if not image_base64:
                        image_base64 = getattr(metadata, 'image', None)
                    
                    if image_base64:
                        # 检测图片格式
                        mime_type = self._detect_image_mime(image_base64)
                        image_info["mime_type"] = mime_type
                        
                        # 保存图片文件
                        if save_images:
                            ext = mime_type.split('/')[-1] if mime_type else 'png'
                            filename = f"{doc_id}-{page_num}-{image_index}.{ext}"
                            file_path = figures_path / filename
                            
                            try:
                                image_data = base64.b64decode(image_base64)
                                with open(file_path, 'wb') as f:
                                    f.write(image_data)
                                image_info["file_path"] = str(file_path)
                                
                                size = self._get_image_size(image_data)
                                if size:
                                    image_info["width"], image_info["height"] = size
                                    
                                logger.debug(f"Saved image: {file_path}")
                            except Exception as e:
                                logger.warning(f"Failed to save image: {e}")
                        
                        # 嵌入 Base64
                        if embed_base64:
                            image_info["base64_data"] = image_base64
                
                # OCR 文字
                ocr_text = getattr(metadata, 'text_as_html', None) or getattr(metadata, 'ocr_text', None)
                if ocr_text:
                    image_info["ocr_text"] = ocr_text
            
            # 如果没有任何有用信息，跳过
            if not any([
                image_info["file_path"],
                image_info["base64_data"],
                image_info["caption"],
                image_info["alt_text"],
                image_info["ocr_text"]
            ]):
                return None
            
            return image_info
            
        except Exception as e:
            logger.warning(f"Failed to process image element: {e}")
            return None
    
    def _process_table_element(
        self,
        element,
        page_num: int,
        table_index: int
    ) -> Optional[Dict[str, Any]]:
        """Process a table element.
        
        优化：将 HTML 表格转换为 Markdown 格式，以便分块器正确解析表头。
        """
        try:
            metadata = element.metadata if hasattr(element, 'metadata') else None
            
            table_info = {
                "page_number": page_num,
                "table_index": table_index,
                "headers": [],
                "rows": [],
                "caption": None,
                "markdown": None,  # 优先使用 HTML 转换的 Markdown
                "html": None
            }
            
            if metadata:
                # 尝试获取结构化表格数据
                text_as_html = getattr(metadata, 'text_as_html', None)
                if text_as_html:
                    table_info["html"] = text_as_html
                    # 将 HTML 转换为 Markdown 表格格式
                    markdown_table = self._html_table_to_markdown(text_as_html)
                    if markdown_table:
                        table_info["markdown"] = markdown_table
                    # 从 HTML 提取表头
                    table_info["headers"] = self._extract_headers_from_html(text_as_html)
            
            # 如果没有成功转换 HTML，降级使用纯文本
            if not table_info["markdown"]:
                table_info["markdown"] = element.text
            
            return table_info
            
        except Exception as e:
            logger.warning(f"Failed to process table element: {e}")
            return None
    
    def _generate_doc_id(self, path: Path) -> str:
        """Generate a short document ID from file path."""
        name = path.stem
        # 取文件名前20个字符 + hash 后4位
        hash_suffix = hashlib.md5(str(path).encode()).hexdigest()[:4]
        return f"{name[:20]}_{hash_suffix}"
    
    def _detect_image_mime(self, base64_data: str) -> str:
        """Detect image MIME type from Base64 data."""
        try:
            # 解码前几个字节来检测格式
            data = base64.b64decode(base64_data[:32] + '==')
            return self._detect_image_mime_from_bytes(data)
        except:
            return 'image/png'
    
    def _detect_image_mime_from_bytes(self, data: bytes) -> str:
        """Detect image MIME type from binary data."""
        try:
            if data[:8] == b'\x89PNG\r\n\x1a\n':
                return 'image/png'
            elif data[:2] == b'\xff\xd8':
                return 'image/jpeg'
            elif data[:6] in (b'GIF87a', b'GIF89a'):
                return 'image/gif'
            elif len(data) >= 12 and data[:4] == b'RIFF' and data[8:12] == b'WEBP':
                return 'image/webp'
            else:
                return 'image/png'  # 默认
        except:
            return 'image/png'
    
    def _get_image_size(self, image_data: bytes) -> Optional[Tuple[int, int]]:
        """Get image dimensions from binary data."""
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(image_data))
            return img.size
        except:
            return None
    
    def _html_table_to_markdown(self, html: str) -> str:
        """将 HTML 表格转换为 Markdown 格式。
        
        Args:
            html: HTML 表格字符串
            
        Returns:
            Markdown 格式的表格字符串
        """
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html, 'html.parser')
            table = soup.find('table')
            if not table:
                return ""
            
            rows = []
            headers = []
            
            # 提取表头 (thead > tr > th 或第一行 tr > th)
            thead = table.find('thead')
            if thead:
                header_row = thead.find('tr')
                if header_row:
                    headers = [th.get_text(strip=True) for th in header_row.find_all(['th', 'td'])]
            
            # 如果没有 thead，检查第一行是否是表头 (tr > th)
            if not headers:
                first_row = table.find('tr')
                if first_row:
                    ths = first_row.find_all('th')
                    if ths:
                        headers = [th.get_text(strip=True) for th in ths]
            
            # 提取数据行 (tbody > tr 或直接 tr)
            tbody = table.find('tbody')
            data_rows = tbody.find_all('tr') if tbody else table.find_all('tr')
            
            for tr in data_rows:
                # 跳过表头行（如果已经处理过）
                if headers and tr.find('th') and not tr.find('td'):
                    continue
                    
                cells = tr.find_all(['td', 'th'])
                row_data = [cell.get_text(strip=True) for cell in cells]
                
                # 如果没有找到表头，使用第一行作为表头
                if not headers and row_data:
                    headers = row_data
                    continue
                
                if row_data:
                    rows.append(row_data)
            
            # 确保所有行的列数一致
            if headers:
                max_cols = max(len(headers), max((len(r) for r in rows), default=0))
                headers = headers + [''] * (max_cols - len(headers))
                rows = [r + [''] * (max_cols - len(r)) for r in rows]
            
            if not headers and not rows:
                return ""
            
            # 构建 Markdown 表格
            md_lines = []
            
            # 表头行
            if headers:
                md_lines.append('| ' + ' | '.join(headers) + ' |')
                # 分隔行
                md_lines.append('|' + '|'.join(['---' for _ in headers]) + '|')
            
            # 数据行
            for row in rows:
                md_lines.append('| ' + ' | '.join(row) + ' |')
            
            return '\n'.join(md_lines)
            
        except ImportError:
            logger.warning("BeautifulSoup not installed, cannot convert HTML table to Markdown")
            return ""
        except Exception as e:
            logger.warning(f"Failed to convert HTML table to Markdown: {e}")
            return ""
    
    def _extract_headers_from_html(self, html: str) -> List[str]:
        """从 HTML 表格中提取表头。
        
        Args:
            html: HTML 表格字符串
            
        Returns:
            表头列表
        """
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html, 'html.parser')
            table = soup.find('table')
            if not table:
                return []
            
            headers = []
            
            # 方式1: 从 thead 提取
            thead = table.find('thead')
            if thead:
                header_row = thead.find('tr')
                if header_row:
                    headers = [th.get_text(strip=True) for th in header_row.find_all(['th', 'td'])]
            
            # 方式2: 从第一行的 th 元素提取
            if not headers:
                first_row = table.find('tr')
                if first_row:
                    ths = first_row.find_all('th')
                    if ths:
                        headers = [th.get_text(strip=True) for th in ths]
            
            # 方式3: 如果没有 th，使用第一行 td 作为表头
            if not headers:
                first_row = table.find('tr')
                if first_row:
                    tds = first_row.find_all('td')
                    if tds:
                        headers = [td.get_text(strip=True) for td in tds]
            
            return headers
            
        except ImportError:
            logger.warning("BeautifulSoup not installed, cannot extract headers from HTML")
            return []
        except Exception as e:
            logger.warning(f"Failed to extract headers from HTML: {e}")
            return []
    
    def _python_docx_table_to_markdown(self, table) -> str:
        """将 python-docx Table 对象转换为 Markdown 格式。
        
        作为 unstructured 表格解析的降级方案。
        
        Args:
            table: python-docx Table 对象
            
        Returns:
            Markdown 格式的表格字符串
        """
        try:
            rows_data = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_cells)
            
            if not rows_data:
                return ""
            
            # 构建 Markdown 表格
            md_lines = []
            
            # 表头（第一行）
            header = rows_data[0]
            md_lines.append("| " + " | ".join(header) + " |")
            
            # 分隔线
            md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            
            # 数据行
            for row in rows_data[1:]:
                # 确保列数一致
                while len(row) < len(header):
                    row.append("")
                md_lines.append("| " + " | ".join(row[:len(header)]) + " |")
            
            return "\n".join(md_lines)
            
        except Exception as e:
            logger.warning(f"Failed to convert python-docx table to Markdown: {e}")
            return ""


# Global instance
unstructured_loader = UnstructuredLoader()
